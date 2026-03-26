"""
Test DOCX conversion to see what's causing the error
"""
import os
import django

# Setup Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'penro_project.settings')
django.setup()

from accounts.models import WorkItemAttachment

# Get attachment ID 16
try:
    attachment = WorkItemAttachment.objects.get(id=16)
    print(f"✓ Found attachment: {attachment.file.name}")
    print(f"  File path: {attachment.file.path}")
    print(f"  File exists: {os.path.exists(attachment.file.path)}")
    print(f"  File size: {attachment.file.size} bytes")
    
    # Try to import docx
    try:
        from docx import Document
        print("✓ python-docx imported successfully")
        
        # Try to load the document
        try:
            doc = Document(attachment.file.path)
            print(f"✓ Document loaded successfully")
            print(f"  Paragraphs: {len(doc.paragraphs)}")
            print(f"  Tables: {len(doc.tables)}")
            
            # Try to convert first paragraph
            if doc.paragraphs:
                print(f"  First paragraph: {doc.paragraphs[0].text[:100]}")
            
        except Exception as e:
            print(f"✗ Error loading document: {e}")
            import traceback
            traceback.print_exc()
            
    except ImportError as e:
        print(f"✗ python-docx not available: {e}")
        
except WorkItemAttachment.DoesNotExist:
    print("✗ Attachment ID 16 not found")
except Exception as e:
    print(f"✗ Error: {e}")
    import traceback
    traceback.print_exc()
